import time
from tkinter import *
from tkinter import filedialog as fd
import openpyxl
import docx
from docx.shared import Pt, Cm
import base64
import requests

# variables
wb_obj = None
excel_path: str

text_convert_done = "Converteren geslaagd!"
text_geen_bestand_gekozen = "Geen bestand gekozen!"
text_bestand_geladen = "Bestand geladen: "

currentVersion = '1.0.2'
url = 'https://api.github.com/repos/Wimpeyy/ToffeToeterTeamnaamTool/contents/version.html'
req = requests.get(url)
if req.status_code == requests.codes.ok:
    req = req.json()  # the response is a JSON
    # req is now a dict with keys: name, encoding, url, size ...
    # and content. But it is encoded with base64.
    content = str(base64.standard_b64decode(req['content']).decode('ascii')).strip()
    if content == currentVersion:
        print("Versie is up to date! Versie: " + content)
    else:
        print("Versieverschil gevonden! Huidige versie: " + currentVersion + ". Nieuwste versie: " + content)
        print("Downloading new version now!")
        newVersion = requests.get("https://github.com/repos/Wimpeyy/ToffeToeterTeamnaamTool/contents/dist/teamnaamtool.exe")
        open("teamnaamtool.exe", "wb").write(newVersion.content)
        print("New version downloaded, restarting in 5 seconds!")
        time.sleep(5)
        quit()
else:
    print('Content was not found.')

# if data == currentVersion:
#     print("App is up to date!")
# else:
#     print("App is not up to date! App is on version " + currentVersion + " but could be on version " + data + "!")
#     print("Downloading new version now!")
#     newVersion = requests.get("https://github.com/yourapp/app-"+data+".exe")
#     open("app.exe", "wb").write(newVersion.content)
#     print("New version downloaded, restarting in 5 seconds!")
#     time.sleep(5)
#     quit()


# Excelfile openen
def open_excel(excelPath):
    global wb_obj
    wb_obj = openpyxl.load_workbook(excelPath)


# command voor converteer button: opslaan naar Wordfile
def save_to_word():
    start_time = textfield_start_time.get()
    if wb_obj is None:
        label_converteren_klaar.config(text=text_geen_bestand_gekozen)
    else:
        word_path = excel_path.rsplit('/', 1)[0]
        bestandsnaam = excel_path.rsplit('/', 1)[1]  # verwijderd alle characters voor de laatste '/'
        word_bestandsnaam = bestandsnaam.split('.', 1)[0]

        sheet_obj = wb_obj.active
        row = sheet_obj.max_row
        doc = docx.Document()

        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)

        # Loop door bestand om cellen te laden en te vullen in Word
        for i in range(1, row + 1):
            cell_teamnaam = sheet_obj.cell(row=i, column=1)
            cell_groepsgrootte = sheet_obj.cell(row=i, column=2)
            paragraph = doc.add_paragraph()
            paragraph.clear()
            run = paragraph.add_run()
            run.font.name = 'Arnhem'

            teamNaam = str(cell_teamnaam.value)
            if len(teamNaam.strip()) <= 24:
                run.font.size = Pt(200)
            else:
                run.font.size = Pt(170)
            # Ruimte genereren tussen teamnamen
            run.add_break()


            # Toevoegen teamnaam
            run = paragraph.add_run()
            run.font.name = "Raleway Black"
            run.font.size = Pt(42)
            run.text = str(cell_teamnaam.value) + " "

            # Line break
            run = paragraph.add_run()
            run.add_break()

            # Toevoegen teamgrootte
            run = paragraph.add_run()
            run.font.name = "Santral-Book"
            run.font.size = Pt(16)
            run.text = str(cell_groepsgrootte.value) + " PERSONEN | " + start_time


            # if i % 2 == 0:
            #     # Forced page break na 2 teams
            #     doc.add_page_break()

        doc.save(word_path + '/' + word_bestandsnaam + ".docx")
        label_converteren_klaar.config(text=text_convert_done + " Bestandsnaam: " + word_bestandsnaam + ".docx")


def select_file():
    filetypes = (
        ('Excel files', '*.xlsx'),
        ('All files', '*.*')
    )

    global excel_path
    excel_path = fd.askopenfilename(
        title='Open a file',
        initialdir='C:/',
        filetypes=filetypes)

    open_excel(excel_path)
    bestandsnaam = excel_path.rsplit('/', 1)[1]  # verwijderd alle characters voor de laatste '/'
    label_bestand_geladen.config(text=text_bestand_geladen + bestandsnaam)


# Window genereren
window = Tk()
window.title('Toffe Toeter Teamnaam Tool')
window.geometry("600x400+10+20")
convertButton = Button(window, text="Converteer", fg='black', command=save_to_word)
convertButton.place(x=500, y=350)

# labels
label_start_time = Label(window, text="Wat is de aanvangtijd?", fg='black', font=("Helvetica", 10))
label_start_time.place(x=77, y=125)
label_converteren_klaar = Label(window, text="", fg='black', font=("Helvetica", 10))
label_converteren_klaar.place(x=100, y=375)
label_bestand_geladen = Label(window, text="", fg='black', font=("Helvetica", 10))
label_bestand_geladen.place(x=100, y=225)

textfield_start_time = Entry(window, bd=2)
textfield_start_time.place(x=80, y=150)

open_excel_button = Button(window, text='Open Excel File', command=select_file)
open_excel_button.place(x=200, y=200)

# run the application
window.mainloop()
